"""
Document Parser Module
Handles parsing of PDF and Word documents to extract text content
"""

import io
import re
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from pathlib import Path
import logging

# PDF parsing libraries
try:
    import pdfplumber
    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    PDF_PLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Word document parsing
try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Data class to hold parsed document information"""
    raw_text: str = ""
    sections: Dict[str, str] = field(default_factory=dict)
    tables: List[List[List[str]]] = field(default_factory=list)
    headings: List[str] = field(default_factory=list)
    bullet_points: List[str] = field(default_factory=list)
    page_count: int = 0
    word_count: int = 0
    file_name: str = ""
    file_type: str = ""
    parse_success: bool = False
    error_message: str = ""


class DocumentParser:
    """
    Main document parser class that handles both PDF and Word documents.
    Uses multiple fallback methods for robust parsing.
    """
    
    # Common SRS section headers
    SECTION_PATTERNS = [
        r'(?i)^[\d.]*\s*(introduction|overview)',
        r'(?i)^[\d.]*\s*(functional\s+requirements?)',
        r'(?i)^[\d.]*\s*(non-?functional\s+requirements?)',
        r'(?i)^[\d.]*\s*(system\s+requirements?)',
        r'(?i)^[\d.]*\s*(user\s+requirements?)',
        r'(?i)^[\d.]*\s*(technical\s+requirements?)',
        r'(?i)^[\d.]*\s*(interface\s+requirements?)',
        r'(?i)^[\d.]*\s*(security\s+requirements?)',
        r'(?i)^[\d.]*\s*(performance\s+requirements?)',
        r'(?i)^[\d.]*\s*(system\s+architecture)',
        r'(?i)^[\d.]*\s*(use\s+cases?)',
        r'(?i)^[\d.]*\s*(features?)',
        r'(?i)^[\d.]*\s*(modules?)',
        r'(?i)^[\d.]*\s*(constraints?)',
        r'(?i)^[\d.]*\s*(assumptions?)',
        r'(?i)^[\d.]*\s*(dependencies?)',
        r'(?i)^[\d.]*\s*(glossary)',
        r'(?i)^[\d.]*\s*(appendix)',
        r'(?i)^[\d.]*\s*(scope)',
        r'(?i)^[\d.]*\s*(objectives?)',
    ]
    
    def __init__(self):
        """Initialize the document parser"""
        self.supported_formats = []
        if PDF_PLUMBER_AVAILABLE or PYPDF2_AVAILABLE:
            self.supported_formats.append('.pdf')
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx', '.doc'])
        
        logger.info(f"DocumentParser initialized. Supported formats: {self.supported_formats}")
    
    def parse(self, file_content: bytes, file_name: str) -> ParsedDocument:
        """
        Main parsing method that routes to appropriate parser based on file type.
        
        Args:
            file_content: Raw bytes of the uploaded file
            file_name: Name of the file including extension
            
        Returns:
            ParsedDocument object containing all extracted information
        """
        result = ParsedDocument(file_name=file_name)
        
        # Determine file type
        file_extension = Path(file_name).suffix.lower()
        result.file_type = file_extension
        
        try:
            if file_extension == '.pdf':
                result = self._parse_pdf(file_content, result)
            elif file_extension in ['.docx', '.doc']:
                result = self._parse_docx(file_content, result)
            else:
                result.error_message = f"Unsupported file format: {file_extension}"
                return result
            
            # Post-processing
            if result.raw_text:
                result = self._post_process(result)
                result.parse_success = True
                
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            result.error_message = f"Error parsing document: {str(e)}"
            result.parse_success = False
            
        return result
    
    def _parse_pdf(self, file_content: bytes, result: ParsedDocument) -> ParsedDocument:
        """
        Parse PDF document using pdfplumber (primary) or PyPDF2 (fallback).
        """
        text_parts = []
        tables = []
        
        # Try pdfplumber first (better table extraction)
        if PDF_PLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    result.page_count = len(pdf.pages)
                    
                    for page in pdf.pages:
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        
                        # Extract tables
                        page_tables = page.extract_tables()
                        if page_tables:
                            tables.extend(page_tables)
                    
                    result.raw_text = '\n\n'.join(text_parts)
                    result.tables = tables
                    logger.info(f"Successfully parsed PDF with pdfplumber: {result.page_count} pages")
                    return result
                    
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                result.page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                result.raw_text = '\n\n'.join(text_parts)
                logger.info(f"Successfully parsed PDF with PyPDF2: {result.page_count} pages")
                return result
                
            except Exception as e:
                raise Exception(f"Both PDF parsers failed: {str(e)}")
        
        raise Exception("No PDF parsing library available")
    
    def _parse_docx(self, file_content: bytes, result: ParsedDocument) -> ParsedDocument:
        """
        Parse Word document using python-docx.
        """
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not available")
        
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            
            text_parts = []
            headings = []
            tables = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
                    
                    # Identify headings
                    if para.style and 'Heading' in para.style.name:
                        headings.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            result.raw_text = '\n\n'.join(text_parts)
            result.headings = headings
            result.tables = tables
            result.page_count = len(doc.paragraphs) // 30 + 1  # Approximate
            
            logger.info(f"Successfully parsed DOCX: {len(text_parts)} paragraphs")
            return result
            
        except PackageNotFoundError:
            raise Exception("Invalid or corrupted Word document")
        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")
    
    def _post_process(self, result: ParsedDocument) -> ParsedDocument:
        """
        Post-process the parsed document to extract structured information.
        """
        text = result.raw_text
        
        # Calculate word count
        result.word_count = len(text.split())
        
        # Extract sections
        result.sections = self._extract_sections(text)
        
        # Extract bullet points
        result.bullet_points = self._extract_bullet_points(text)
        
        # Extract additional headings from text if not already found
        if not result.headings:
            result.headings = self._extract_headings(text)
        
        return result
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract document sections based on common SRS headers.
        """
        sections = {}
        lines = text.split('\n')
        current_section = "Introduction"
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            is_header = False
            for pattern in self.SECTION_PATTERNS:
                match = re.match(pattern, line)
                if match:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = line
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_bullet_points(self, text: str) -> List[str]:
        """
        Extract bullet points and numbered lists from text.
        """
        bullet_patterns = [
            r'^\s*[•●○◦▪▫►]\s*(.+)',  # Bullet characters
            r'^\s*[-*]\s+(.+)',  # Dash or asterisk bullets
            r'^\s*\d+[.)]\s+(.+)',  # Numbered lists
            r'^\s*[a-zA-Z][.)]\s+(.+)',  # Lettered lists
            r'^\s*(?:FR|NFR|UC|REQ)[-_]?\d+[.:]\s*(.+)',  # Requirement IDs
        ]
        
        bullet_points = []
        for line in text.split('\n'):
            for pattern in bullet_patterns:
                match = re.match(pattern, line)
                if match:
                    bullet_points.append(match.group(1).strip())
                    break
        
        return bullet_points
    
    def _extract_headings(self, text: str) -> List[str]:
        """
        Extract potential headings from text based on patterns.
        """
        headings = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for heading patterns
            # Pattern 1: All caps (likely heading)
            if line.isupper() and len(line.split()) <= 7:
                headings.append(line)
                continue
            
            # Pattern 2: Numbered sections (1. or 1.1 or 1.1.1)
            if re.match(r'^[\d.]+\s+\w+', line) and len(line.split()) <= 10:
                headings.append(line)
                continue
            
            # Pattern 3: Title case followed by colon
            if re.match(r'^[A-Z][a-zA-Z\s]+:\s*$', line):
                headings.append(line.rstrip(':').strip())
        
        return headings
    
    def get_supported_formats(self) -> List[str]:
        """Return list of supported file formats"""
        return self.supported_formats.copy()


class DocumentParserFactory:
    """
    Factory class for creating document parsers.
    Allows for future extensibility with different parser implementations.
    """
    
    _instance = None
    
    @classmethod
    def get_parser(cls) -> DocumentParser:
        """Get or create a singleton parser instance"""
        if cls._instance is None:
            cls._instance = DocumentParser()
        return cls._instance


def parse_document(file_content: bytes, file_name: str) -> ParsedDocument:
    """
    Convenience function for parsing documents.
    
    Args:
        file_content: Raw bytes of the uploaded file
        file_name: Name of the file including extension
        
    Returns:
        ParsedDocument object containing all extracted information
    """
    parser = DocumentParserFactory.get_parser()
    return parser.parse(file_content, file_name)